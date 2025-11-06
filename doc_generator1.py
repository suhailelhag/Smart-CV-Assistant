# doc_generator1.py

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ملاحظة: مستويات الإتقان محفوظة بالإنجليزية في البيانات

# +++ دوال مساعدة للتصميم الجديد +++

def add_horizontal_line(paragraph):
    """
    تضيف خطًا أفقيًا أسفل الفقرة عن طريق تعيين حدودها السفلية.
    """
    p_border = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '6')      # حجم الخط
    bottom_border.set(qn('w:space'), '1')   # المسافة من النص
    bottom_border.set(qn('w:color'), 'auto') # اللون
    p_border.append(bottom_border)
    paragraph._p.get_or_add_pPr().append(p_border)

def add_section_header(doc, title):
    """
    تضيف عنوان قسم منسق مع خط أفقي تحته.
    """
    # أضف مسافة قبل العنوان إذا لم يكن العنصر الأول في المستند
    if len(doc.paragraphs) > 1:
        doc.add_paragraph().paragraph_format.space_before = Pt(12)
    
    p = doc.add_paragraph()
    run = p.add_run(title.upper()) # تحويل العنوان إلى أحرف كبيرة
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(12)
    font.bold = True
    
    add_horizontal_line(p)
    p.paragraph_format.space_after = Pt(8)

def write_cv(user_data: dict, ai_data: dict, file_path: str, section_names: dict, section_order: list):
    """
    تكتب بيانات المستخدم والبيانات التي تم إنشاؤها بواسطة الذكاء الاصطناعي في ملف وورد
    باستخدام التصميم الجديد، بناءً على الترتيب المخصص من قبل المستخدم.
    """
    doc = Document()

    # --- إعدادات المستند العامة (التصميم الجديد) ---
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- قسم الرأس (التصميم الجديد) ---
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(user_data.get('name', ''))
    run_name.font.size = Pt(22)
    run_name.bold = True
    p_name.paragraph_format.space_after = Pt(2)
    
    # معلومات الاتصال
    contact_info = []
    if user_data.get('linkedin'): contact_info.append(user_data['linkedin'])
    if user_data.get('email'): contact_info.append(user_data['email'])
    if user_data.get('location'): contact_info.append(user_data['location'])
    if user_data.get('phone'): contact_info.append(user_data['phone'])
    
    p_contact = doc.add_paragraph(' | '.join(contact_info))
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.runs[0].font.size = Pt(11)
    p_contact.paragraph_format.space_after = Pt(12)

    # --- الحلقة الرئيسية لكتابة الأقسام (باستخدام التصميم الجديد) ---
    for section_key in section_order:
        # -- قسم الملخص المهني --
        if section_key == 'profile' and ai_data.get('profile'):
            add_section_header(doc, section_names.get('profile', 'Professional Summary'))
            p_profile = doc.add_paragraph(ai_data['profile'])
            p_profile.paragraph_format.space_after = Pt(4)
        
        # -- قسم الخبرة العملية --
        elif section_key == 'experiences' and user_data.get('experiences'):
            add_section_header(doc, section_names.get('experiences', 'Experience'))
            for exp in user_data['experiences']:
                # السطر الأول: المسمى الوظيفي | اسم الشركة
                p_job_title = doc.add_paragraph()
                run_title = p_job_title.add_run(exp.get('position', ''))
                run_title.font.bold = True
                p_job_title.add_run(f" | {exp.get('company', '')}")
                p_job_title.paragraph_format.space_after = Pt(0)
                
                # السطر الثاني: الموقع      <tab>      المدة
                p_loc_date = doc.add_paragraph()
                p_loc_date.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
                p_loc_date.add_run(exp.get('location', ''))
                p_loc_date.add_run('\t')
                run_duration = p_loc_date.add_run(exp.get('duration', ''))
                run_duration.font.italic = True
                p_loc_date.paragraph_format.space_after = Pt(4)

                # الإنجازات والتفاصيل
                for detail in exp.get('details', []):
                    doc.add_paragraph(detail, style='List Bullet')
                    doc.paragraphs[-1].paragraph_format.space_after = Pt(2)
                
                # إضافة مسافة بعد آخر نقطة في كل خبرة
                if doc.paragraphs[-1].style.name.startswith('List'):
                    doc.paragraphs[-1].paragraph_format.space_after = Pt(10)

        # -- قسم التعليم --
        elif section_key == 'education' and user_data.get('degree'):
            add_section_header(doc, section_names.get('education', 'Education'))
            p_university = doc.add_paragraph()
            run_university = p_university.add_run(user_data.get('university', ''))
            run_university.font.bold = True
            p_university.paragraph_format.space_after = Pt(0)
            
            p_degree = doc.add_paragraph(user_data.get('degree', ''))
            p_degree.paragraph_format.space_after = Pt(8)

        # -- قسم الشهادات --
        elif section_key == 'certifications' and user_data.get('certifications'):
            add_section_header(doc, section_names.get('certifications', 'Certifications'))
            for cert in user_data['certifications']:
                p_cert = doc.add_paragraph()
                run_cert_name = p_cert.add_run(cert.get('name', ''))
                run_cert_name.font.bold = True
                p_cert.paragraph_format.space_after = Pt(0)
                
                p_authority = doc.add_paragraph(cert.get('authority', ''))
                p_authority.paragraph_format.space_after = Pt(6)

        # -- قسم اللغات --
        elif section_key == 'languages' and user_data.get('languages'):
            add_section_header(doc, section_names.get('languages', 'Languages'))
            for lang in user_data['languages']:
                p_lang = doc.add_paragraph()
                run_lang_name = p_lang.add_run(f"{lang.get('name', '')}: ")
                run_lang_name.font.bold = True
                p_lang.add_run(lang.get('proficiency', ''))
                p_lang.paragraph_format.space_after = Pt(2)

        # -- الأقسام الديناميكية الأخرى (مثل المهارات والاهتمامات) --
        elif section_key in ai_data and section_key in section_names and ai_data[section_key]:
            add_section_header(doc, section_names[section_key])
            items_list = ai_data[section_key]
            
            # عرض المهارات أو أي قوائم أخرى كنقاط
            for item in items_list:
                doc.add_paragraph(item, style='List Bullet')
                doc.paragraphs[-1].paragraph_format.space_after = Pt(2)

    doc.save(file_path)
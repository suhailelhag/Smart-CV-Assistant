# doc_generator.py

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ملاحظة: مستويات الإتقان محفوظة بالإنجليزية في البيانات

# +++ Funciones de Ayuda del Diseño Moderno +++
def set_run_font(run, name='Calibri', size_pt=10.5, bold=False, italic=False, color_rgb=None):
    """Función de ayuda para establecer las propiedades de fuente de un 'run'."""
    font = run.font
    font.name = name
    font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    if color_rgb:
        font.color.rgb = RGBColor.from_string(color_rgb)

def add_new_section_title(doc, title):
    """
    Agrega un título flanqueado por líneas horizontales usando una tabla de tres columnas.
    """
    # Evita agregar espacio extra si el título es el primer elemento
    if len(doc.paragraphs) > 1:
        doc.add_paragraph()

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Ancho de las columnas ajustado para un diseño equilibrado
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(2.5) # Ancho aumentado para títulos más largos
    table.columns[2].width = Inches(2.5)

    # Celda izquierda con la línea
    cell_left = table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    p_left.add_run("__________________________________")
    p_left.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Celda central con el título
    cell_title = table.cell(0, 1)
    p_title = cell_title.paragraphs[0]
    run_title = p_title.add_run(title.upper())
    set_run_font(run_title, name='Calibri', size_pt=12, bold=True)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Celda derecha con la línea
    cell_right = table.cell(0, 2)
    p_right = cell_right.paragraphs[0]
    p_right.add_run("__________________________________")
    p_right.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Asegura que el contenido de la celda esté centrado verticalmente y sin espaciado
    for cell in table.rows[0].cells:
        cell.vertical_alignment = 1 # 1 para centrar
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            
    doc.add_paragraph().paragraph_format.space_before = Pt(8)


def write_cv(user_data: dict, ai_data: dict, file_path: str, section_names: dict, section_order: list):
    """
    Escribe los datos del usuario y los generados por la IA en un archivo Word
    utilizando un diseño moderno, basado en el orden personalizado del usuario.
    """
    doc = Document()

    # --- Configuración General del Documento (Diseño Moderno) ---
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)

    # --- Sección de Encabezado (Diseño Moderno) ---
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(user_data.get('name', ''))
    set_run_font(run_name, name='Calibri', size_pt=26, bold=True)
    p_name.paragraph_format.space_after = Pt(4)
    
    # Título del puesto debajo del nombre (del script original)
    if user_data.get('title'):
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run(user_data['title'])
        set_run_font(run_title, size_pt=14)
        p_title.paragraph_format.space_after = Pt(6)

    contact_info = []
    if user_data.get('location'): contact_info.append(user_data['location'])
    if user_data.get('phone'): contact_info.append(user_data['phone'])
    if user_data.get('email'): contact_info.append(user_data['email'])
    if user_data.get('linkedin'): contact_info.append(user_data['linkedin'])
    
    p_contact = doc.add_paragraph('  ♦  '.join(contact_info))
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(6)

    p_header_line = doc.add_paragraph("_________________________________________________________________")
    p_header_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header_line.paragraph_format.space_after = Pt(8)

    # --- Bucle Principal para Escribir Secciones (Lógica Original con Diseño Moderno) ---
    for section_key in section_order:
        # -- Sección de Resumen Profesional --
        if section_key == 'profile' and ai_data.get('profile'):
            add_new_section_title(doc, section_names.get('profile', 'Professional Summary'))
            p_profile = doc.add_paragraph(ai_data['profile'])
            p_profile.paragraph_format.space_after = Pt(4)
        
        # -- Sección de Experiencia Laboral --
        elif section_key == 'experiences' and user_data.get('experiences'):
            add_new_section_title(doc, section_names.get('experiences', 'Experience'))
            for exp in user_data['experiences']:
                p_job_title = doc.add_paragraph()
                run_title = p_job_title.add_run(exp.get('position', ''))
                set_run_font(run_title, bold=True)
                
                # Poner la duración a la derecha
                p_job_title.add_run('\t')
                run_duration = p_job_title.add_run(exp.get('duration', ''))
                set_run_font(run_duration, italic=True)
                
                # Configurar tabulación a la derecha
                tab_stops = p_job_title.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
                
                p_job_title.paragraph_format.space_after = Pt(0)
                
                p_company = doc.add_paragraph()
                run_company = p_company.add_run(exp.get('company', ''))
                set_run_font(run_company, italic=True)
                p_company.paragraph_format.space_after = Pt(4)
                
                for detail in exp.get('details', []):
                    doc.add_paragraph(detail, style='List Bullet')
                    doc.paragraphs[-1].paragraph_format.space_after = Pt(2)
                
                # Añadir más espacio después de la última viñeta de cada experiencia
                if doc.paragraphs[-1].style.name.startswith('List'):
                    doc.paragraphs[-1].paragraph_format.space_after = Pt(12)

        # -- Sección de Educación --
        elif section_key == 'education' and user_data.get('degree'):
            add_new_section_title(doc, section_names.get('education', 'Education'))
            p_education_title = doc.add_paragraph()
            run_degree = p_education_title.add_run(user_data.get('degree', ''))
            set_run_font(run_degree, bold=True)
            p_education_title.paragraph_format.space_after = Pt(0)
            
            p_university = doc.add_paragraph(user_data.get('university', ''))
            p_university.paragraph_format.space_after = Pt(8)

        # -- Sección de Certificaciones --
        elif section_key == 'certifications' and user_data.get('certifications'):
            add_new_section_title(doc, section_names.get('certifications', 'Certifications'))
            for cert in user_data['certifications']:
                p_cert = doc.add_paragraph()
                run_cert_name = p_cert.add_run(cert.get('name', ''))
                set_run_font(run_cert_name, bold=True)
                p_cert.add_run(f" - {cert.get('authority', '')}")
                p_cert.paragraph_format.space_after = Pt(4)

        # -- Sección de Idiomas --
        elif section_key == 'languages' and user_data.get('languages'):
            add_new_section_title(doc, section_names.get('languages', 'Languages'))
            for lang in user_data['languages']:
                p_lang = doc.add_paragraph()
                run_lang_name = p_lang.add_run(f"{lang.get('name', '')}: ")
                set_run_font(run_lang_name, bold=True)
                p_lang.add_run(lang.get('proficiency', '')) # Usar el nivel de competencia directamente
                p_lang.paragraph_format.space_after = Pt(2)
        
        # -- Otras Secciones Dinámicas (como Habilidades, Intereses) --
        elif section_key in ai_data and section_key in section_names and ai_data[section_key]:
            add_new_section_title(doc, section_names[section_key])
            items_list = ai_data[section_key]
            
            # Caso especial para 'skills' para usar una tabla de 2 columnas
            if section_key == 'skills':
                # Asegura que la lista no esté vacía
                if not items_list: continue
                num_rows = (len(items_list) + 1) // 2
                table = doc.add_table(rows=num_rows, cols=2)
                table.autofit = False
                table.columns[0].width = Inches(3.75)
                table.columns[1].width = Inches(3.75)
                
                for i in range(num_rows):
                    if (i*2) < len(items_list): table.cell(i, 0).text = f"•  {items_list[i*2]}"
                    if (i*2+1) < len(items_list): table.cell(i, 1).text = f"•  {items_list[i*2+1]}"
                
                for row in table.rows:
                    for cell in row.cells: cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            else:
                for item in items_list:
                    doc.add_paragraph(item, style='List Bullet')
                    doc.paragraphs[-1].paragraph_format.space_after = Pt(2)

    doc.save(file_path)